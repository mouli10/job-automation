import json
import logging
from pathlib import Path
import docx
from docx.text.paragraph import Paragraph
from google import genai
from src.config import GEMINI_API_KEY, admin_config, DATA_DIR

logger = logging.getLogger(__name__)

GEMINI_MODEL = "gemini-2.5-flash-lite"
OPTIMIZED_RESUMES_DIR = DATA_DIR / "resumes" / "optimized"
OPTIMIZED_RESUMES_DIR.mkdir(parents=True, exist_ok=True)

OPTIMIZER_PROMPT = """You are an elite Resume Rewriter. Your ONLY task is to subtly inject important keywords from the Job Description into the applicant's existing bullet points without changing their core meaning.

CRITICAL SAFETY & QUALITY CONSTRAINTS:
1. NEVER remove or generalize numerical metrics, quantifiable impact, or measurable achievements (DO NOT change "~6 hours to under 1 hour" into "reduced time").
2. PREFER MINIMAL EDITS. If a bullet point is already strong or irrelevant to the JD, return the original text untouched.
3. DO NOT fabricate or infer experience. NEVER add "years of experience" or subjective proficiency ("2-3 years", "expert level").
4. DO NOT add generic trailing phrases (e.g., "supporting business intelligence", "for enterprise environments", "for business and operations teams", "supporting data-driven decision making", "supporting enterprise data solutions").
5. AVOID Repetitive Phrasing. Ensure your edits read like human-edited content, not templated AI injections repeating the same phrase across multiple bullets.
6. Prefer additive improvements (e.g. gently inserting relevant tool keywords into the middle of the sentence) rather than rewriting the core structure.
7. DO NOT add new achievements. Only optimize wording within existing achievements.
8. Do NOT drastically change the length of the bullet point.
9. Do NOT use markdown formatting like **bold** or *italics*. Return pure plain text.

Job Description Snippet:
{job_desc}

Bullet Points to Optimize:
{bullets_json}

Return ONLY a valid JSON object matching this schema exactly:
{{"changes": [{{"original": "exact original text", "optimized": "new text"}}]}}
IMPORTANT: Do NOT output markdown ticks (```json). Return ONLY the raw JSON object.
"""

class ResumeOptimizer:
    def __init__(self):
        try:
            self.client = genai.Client(api_key=GEMINI_API_KEY)
        except Exception as e:
            logger.error(f"Failed to initialize Gemini Client for Optimizer: {e}")
            self.client = None

    def optimize(self, doc_path: str, job_id: int, job_title: str, job_desc: str) -> str:
        """Runs the optimization pipeline safely on a docx file."""
        if not self.client:
            return ""
            
        doc = docx.Document(doc_path)
        bullets = self._extract_bullets(doc)
        
        if not bullets:
            logger.warning(f"  ⚠️  No extractable bullet points found in {Path(doc_path).name}")
            return ""
            
        logger.info(f"  🧠 Sending {len(bullets)} bullets to Gemini for optimization...")
        
        changes = self._call_gemini_optimize(bullets, job_desc)
        if not changes:
            logger.warning("  ⚠️  Gemini failed to return valid optimization changes.")
            return ""

        replacements_made = self._apply_safely(doc, changes)
        if replacements_made == 0:
            logger.warning("  ⚠️  Failed to match and replace any optimized bullets.")
            return ""

        out_name = f"{Path(doc_path).stem}_opt_{job_id}.docx"
        out_path = OPTIMIZED_RESUMES_DIR / out_name
        doc.save(str(out_path))
        logger.info(f"  ✅ Saved optimized resume: {out_name} with {replacements_made} edits.")
        return str(out_path)

    def _iter_paragraphs(self, doc):
        """Yields all paragraphs in the document, including those inside tables."""
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    def _extract_bullets(self, doc) -> list:
        targets = []
        skip_section = False
        
        for p in self._iter_paragraphs(doc):
            text = p.text.strip()
            
            # Detect Sections to skip (e.g. Education)
            upper_text = text.upper()
            if "EDUCATION" in upper_text and len(upper_text) < 20:
                skip_section = True
                continue
            elif skip_section and len(upper_text) > 0 and len(upper_text) < 30 and (
                "EXPERIENCE" in upper_text or "PROJECTS" in upper_text or "SKILLS" in upper_text):
                # If we hit another major header, stop skipping
                skip_section = False

            if skip_section:
                continue

            text = text.replace("**", "").replace("*", "")
            # Most actual bullets or experience points are longer than 40 characters
            if len(text) > 40:
                targets.append(text)
        return targets

    def _call_gemini_optimize(self, bullets: list, job_desc: str) -> list:
        try:
            prompt = OPTIMIZER_PROMPT.format(
                job_desc=job_desc[:3000],
                bullets_json=json.dumps(bullets)
            )
            response = self.client.models.generate_content(model=GEMINI_MODEL, contents=prompt)
            text = response.text.replace("```json", "").replace("```", "").strip()
            data = json.loads(text)
            return data.get("changes", [])
        except Exception as e:
            logger.error(f"Error in Gemini optimizer call: {e}")
            return []

    def _apply_safely(self, doc, changes: list) -> int:
        """Replaces paragraph text inline while maintaining first run formatting completely."""
        replaced_count = 0
        
        # Build mapping for O(1) lookup
        mapping = {}
        for c in changes:
            orig = c.get("original", "").strip().replace("**", "").replace("*", "")
            opt = c.get("optimized", "").strip().replace("**", "").replace("*", "")
            if orig and opt and orig != opt:
                mapping[orig] = opt

        for p in self._iter_paragraphs(doc):
            original_text = p.text.strip().replace("**", "").replace("*", "")
            if original_text in mapping:
                new_text = mapping[original_text]
                
                # Capture baseline font from the first available run
                base_name = None
                base_size = None
                base_bold = None
                base_italic = None
                
                if p.runs:
                    run = p.runs[0]
                    base_name = run.font.name
                    base_size = run.font.size
                    base_bold = run.font.bold
                    base_italic = run.font.italic
                
                # Clear all existing runs inside paragraph
                p.clear()
                
                # Append new text as a single run with the baseline style
                new_run = p.add_run(new_text)
                if base_name: new_run.font.name = base_name
                if base_size: new_run.font.size = base_size
                if base_bold is not None: new_run.font.bold = base_bold
                if base_italic is not None: new_run.font.italic = base_italic

                replaced_count += 1
                
        return replaced_count
